import asyncio
import os
from pathlib import Path
import pandas as pd
from pyannote.audio import Pipeline
import whisperx
import pydub
import torch
import requests
import json
from datetime import timedelta
import subprocess
import app.requests as rq
import pandas as pd
from datetime import timedelta
from docx import *
import config


# async def to_wav(audio_path):
#     filename = Path(audio_path)
#     filename = filename.stem
#     subprocess.call(['ffmpeg', '-i', f'{audio_path}',
#                    '/home/jupyter/datasphere/project/test4.wav'])

async def load_pipeline_from_pretrained(path_to_config: str | Path) -> Pipeline:
    path_to_config = Path(path_to_config)
    print(f"Loading pyannote pipeline from {path_to_config}...")
    pipeline = Pipeline.from_pretrained(path_to_config)
    return pipeline

async def diarize_transcript_audio(audio_path, otchet_id):
    # PATH_TO_CONFIG = config.PATH_TO_CONFIG
    # pipeline = await load_pipeline_from_pretrained(PATH_TO_CONFIG)
    # device = config.device
    # pipeline = pipeline.to(torch.device(device))
    # whisper_model = whisperx.load_model("small", device, compute_type="int8")
    # audio_file = audio_path
    # filename = Path(audio_path)
    # filename = filename.stem

    # # Выполнение диаризации на всей записи
    # print("Диаризирую")
    # diarization_result = pipeline({"audio": audio_file})

    # # Преобразование результата диаризации в DataFrame
    # diarization_list = []
    # for segment, track, label in diarization_result.itertracks(yield_label=True):
    #     diarization_list.append({
    #         'start': segment.start,
    #         'end': segment.end,
    #         'speaker': label
    #     })

    # diarization_df = pd.DataFrame(diarization_list)
    # print("Diarization DataFrame:")
    # print(diarization_df)

    # # Транскрипция всей записи
    # print("Транскрибирую")
    # transcription_result = whisper_model.transcribe(audio_file)
    # print("Закончил транскрипцию")

    # # Выравнивание сегментов
    # aligned_model, metadata = whisperx.load_align_model(
    #     language_code=transcription_result["language"], device=device
    # )
    # aligned_result = whisperx.align(
    #     transcription_result["segments"], aligned_model, metadata, audio_file, device
    # )

    # # Присвоение говорящих сегментам
    # segments_with_speakers = whisperx.assign_word_speakers(diarization_df, aligned_result)

    # # Преобразование итогового результата транскрипции с говорящими в DataFrame
    # all_transcription_list = []
    # for segment in segments_with_speakers['segments']:
    #     if isinstance(segment, dict):
    #         speaker = segment.get('speaker', 'unknown')
    #         text = segment.get('text', '')
    #         start = segment.get('start', 0)
    #         end = segment.get('end', 0)

    #         all_transcription_list.append({
    #             'speaker': speaker,
    #             'text': text,
    #             'start': start,
    #             'end': end
    #         })

    # transcription_df = pd.DataFrame(all_transcription_list)
    # print("Итоговая транскрипция с говорящими:")
    # print(transcription_df)

    # # Подсчет уникальных спикеров
    # unique_speakers = diarization_df['speaker'].unique()
    # num_speakers = len(unique_speakers)
    # print(f"Количество уникальных спикеров: {num_speakers}")


    # transcription_df.to_csv(f"app/outputs/{filename}/transcription.csv",sep=";")
    # rq.add_output_csv_path(otchet_id, f"app/outputs/{filename}/transcription.csv")
    # 
    pass

async def get_text_with_time(audio_path):
    filename = Path(audio_path)
    filename = filename.stem

    df = pd.read_csv(f"app/outputs/{filename}/transcription.csv", sep=';', encoding='utf-8')
    def format_time(seconds):
        td = timedelta(seconds=seconds)
        minutes, seconds = divmod(td.total_seconds(), 60)
        return f"{int(minutes):02}:{int(seconds):02}"

    # Примените форматирование к каждой строке и сохраните результат
    formatted_texts = []
    for _, row in df.iterrows():
        start_time = format_time(row['start'])
        end_time = format_time(row['end'])
        formatted_texts.append(f"({start_time} - {end_time}) {row['text']}")

    # Выводите или сохраняйте результат
    full_text = ""

    for line in formatted_texts:
        full_text = full_text + line + "\n"
        
    return full_text
async def get_text(audio_path):
    filename = Path(audio_path)
    filename = filename.stem
    df = pd.read_csv(f"app/outputs/{filename}/transcription.csv", sep=';', encoding='utf-8')

    formatted_texts = []
    for _, row in df.iterrows():
        formatted_texts.append(f"{row['text']}")
    full_text = ""
    for line in formatted_texts:
        full_text = full_text + line + "\n"
        
    return full_text

async def yandexgpt(text, prompt):
    token = config.token
    url = 'https://llm.api.cloud.yandex.net/foundationModels/v1/completion'
    headers = {
        'Authorization': f'Bearer {token}',
        'Content-Type': 'application/json'
    }
    data = {
        "modelUri": "gpt://b1g72uajlds114mlufqi/yandexgpt/latest",
        "completionOptions": {
            "stream": False,
            "temperature": 0.6,
            "maxTokens": 2000
        },
        "messages": [
            {
                "role": "system",
                "text": prompt
            },
            {
                "role": "user",
                "text": text
            }
        ]
    }

    response = requests.post(url, headers=headers, json=data) 
    if response.status_code == 200:
        # print(response.json())
        json_dump = json.dumps(response.json(), ensure_ascii=False)
        answer = json.loads(json_dump)['result']['alternatives'][0]['message']['text']
        return answer
    else:
        print(f"Ошибка: {response.status_code}")
        return "Error"

async def format_time(seconds):
    td = timedelta(seconds=seconds)
    minutes, seconds = divmod(td.total_seconds(), 60)
    return f"{int(minutes):02}:{int(seconds):02}"

async def create_descrypt_file(audio_path):
    filename = Path(audio_path)
    filename = filename.stem
    df = pd.read_csv(f"app/outputs/{filename}/transcription.csv", sep=';', encoding='utf-8')

    doc = Document("app/decryption.docx")
    formatted_texts = []
    current_speaker_id = -1
    for _, row in df.iterrows():
        speaker_id = int(row['speaker'].replace("SPEAKER_0", "")) + 1
        speaker = f"Спикер {speaker_id}:"
        start_time = await format_time(row['start'])
        end_time = await format_time(row['end'])
        if speaker_id != current_speaker_id:
            formatted_texts.append(f"{speaker}\n({start_time} - {end_time}) {row['text']}")
            current_speaker_id = speaker_id
        else:
            formatted_texts.append(f"({start_time} - {end_time}) {row['text']}")

    for line in formatted_texts:
        doc.add_paragraph(line)
    doc.save(f"app/outputs/{filename}/decryption.docx")

async def create_doc_from_text(text, audio_path):
    # Создаем документ
    doc = Document()
    filename = Path(audio_path)
    filename = filename.stem

    # Разбиваем текст на строки
    lines = text.strip().split("\n")

    # Пройдем по строкам и добавим форматирование
    for line in lines:
        line = line.strip()
        if line.startswith("ТЕМА ВСТРЕЧИ:") or line.startswith("**Тема встречи:**"):
            # Добавляем заголовок темы встречи
            doc.add_heading(line, level=1)
        elif line.startswith("ПРОТОКОЛ ВСТРЕЧИ") or line.startswith("**Протокол встречи:**"):
            doc.add_heading("Протокол встречи", level=2)
        elif line.startswith("Дата:") or line.startswith("* **Дата:**"):
            # Добавляем дату
            doc.add_paragraph(line, style='Normal')
        elif line.startswith("Время:") or line.startswith("* **Время начала и окончания совещания:**"):
            # Добавляем время
            doc.add_paragraph(line, style='Normal')
        elif line.startswith("Длительность:") or line.startswith("* **Длительность:**"):
            # Добавляем длительность
            doc.add_paragraph(line, style='Normal')
        elif line.startswith("Участники:") or line.startswith("* **Участники:**"):
            doc.add_heading("Участники", level=3)
        elif line.startswith("* ") or line.startswith("    *"):
            # Добавляем список участников
            doc.add_paragraph(line.replace("* ", "").replace("    *", ""), style='List Bullet')
        elif line.startswith("ПОВЕСТКА ДНЯ") or line.startswith("**Повестка дня:**"):
            doc.add_heading("Повестка дня", level=2)
        elif line.startswith("Цель встречи:") or line.startswith("* краткое описание цели встречи:"):
            # Добавляем описание цели встречи
            doc.add_paragraph(line, style='Normal')
        elif line.startswith("РЕЗЮМЕ ОБСУЖДЕНИЙ") or line.startswith("**I. Название вопроса**:"):
            doc.add_heading("Резюме обсуждений", level=2)
        elif line.startswith("Решено:") or line.startswith("* **Решено:**"):
            # Добавляем решение
            doc.add_paragraph(line, style='Normal')
        else:
            doc.add_paragraph(line, style='Normal')

    # Сохраняем документ
    doc.save(f"app/outputs/{filename}/docx.docx")


if __name__ == "__main__":
    asyncio.run(diarize_transcript_audio("test2.wav"))